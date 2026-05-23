from __future__ import annotations

import io
import logging
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from fastapi import HTTPException, UploadFile

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
ALLOWED_MIME_TYPES = {
    ".txt": {"text/plain"},
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}
MAX_PDF_PAGES = 50
MAX_EXTRACTED_CHARS = 200_000
MAX_DOCX_FILES = 1_000
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
PREVIEW_LEN = 500
READ_CHUNK_BYTES = 1024 * 1024
logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    filename: str
    file_type: str


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def _parse_txt(data: bytes) -> str:
    if b"\x00" in data[:1024]:
        raise ValueError("텍스트 파일에 바이너리 데이터가 포함되어 있습니다.")
    for encoding in ("utf-8", "utf-8-sig", "cp949", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("텍스트 파일 인코딩을 해석할 수 없습니다.")


def _parse_pdf(data: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("PDF 파서가 서버에 설치되어 있지 않습니다.") from exc

    parts: list[str] = []
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            if doc.page_count > MAX_PDF_PAGES:
                raise ValueError(f"PDF는 최대 {MAX_PDF_PAGES}페이지까지 업로드할 수 있습니다.")
            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    parts.append(page_text)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("PDF 파일을 읽을 수 없습니다. 손상되었거나 암호화된 파일일 수 있습니다.") from exc
    if not parts:
        raise ValueError("PDF에서 추출할 텍스트가 없습니다.")
    return "\n\n".join(parts)


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX 파서가 서버에 설치되어 있지 않습니다.") from exc

    _validate_docx_archive(data)
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError("Word 문서를 읽을 수 없습니다. 손상되었거나 지원하지 않는 DOCX 파일입니다.") from exc
    parts: list[str] = []

    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise ValueError("Word 문서에서 추출할 텍스트가 없습니다.")
    return "\n".join(parts)


def _validate_mime_type(ext: str, content_type: str | None) -> None:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    if not normalized or normalized == "application/octet-stream":
        return
    allowed = ALLOWED_MIME_TYPES.get(ext, set())
    if normalized not in allowed:
        raise HTTPException(
            status_code=415,
            detail=f"파일 MIME 타입이 올바르지 않습니다. 감지: {content_type}, 허용: {', '.join(sorted(allowed))}",
        )


def _validate_magic_bytes(ext: str, data: bytes) -> None:
    if ext == ".pdf" and not data.startswith(b"%PDF-"):
        raise HTTPException(status_code=415, detail="PDF 파일 서명이 올바르지 않습니다.")
    if ext == ".docx" and not data.startswith(b"PK\x03\x04"):
        raise HTTPException(status_code=415, detail="DOCX 파일 서명이 올바르지 않습니다.")
    if ext == ".txt" and b"\x00" in data[:1024]:
        raise HTTPException(status_code=415, detail="TXT 파일에 바이너리 데이터가 포함되어 있습니다.")


def _validate_pdf_page_hint(data: bytes) -> None:
    # Full page validation happens in PyMuPDF, but this cheap hint rejects obviously huge PDFs early.
    page_markers = re.findall(rb"/Type\s*/Page\b", data)
    if len(page_markers) > MAX_PDF_PAGES:
        raise HTTPException(
            status_code=413,
            detail=f"PDF는 최대 {MAX_PDF_PAGES}페이지까지 업로드할 수 있습니다.",
        )


def _validate_docx_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOCX_FILES:
                raise ValueError("DOCX 내부 파일 수가 너무 많습니다.")
            total_size = sum(info.file_size for info in infos)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("DOCX 압축 해제 크기가 너무 큽니다.")
            if "word/document.xml" not in {info.filename for info in infos}:
                raise ValueError("유효한 Word 문서 구조가 아닙니다.")
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX 압축 구조가 올바르지 않습니다.") from exc


def parse_bytes(filename: str, data: bytes, content_type: str | None = None) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"지원하지 않는 형식입니다. 허용: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    _validate_mime_type(ext, content_type)
    _validate_magic_bytes(ext, data)
    if ext == ".pdf":
        _validate_pdf_page_hint(data)

    if ext == ".txt":
        try:
            raw = _parse_txt(data)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
    elif ext == ".pdf":
        try:
            raw = _parse_pdf(data)
        except ValueError as exc:
            status_code = 413 if "최대" in str(exc) else 422
            raise HTTPException(status_code=status_code, detail=str(exc)) from exc
    else:
        try:
            raw = _parse_docx(data)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc

    text = _normalize_text(raw)
    if not text:
        raise HTTPException(status_code=422, detail="문서에서 유효한 텍스트를 추출하지 못했습니다.")
    if len(text) > MAX_EXTRACTED_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"추출된 문서 텍스트는 최대 {MAX_EXTRACTED_CHARS:,}자까지 처리할 수 있습니다.",
        )

    return ParsedDocument(text=text, filename=filename, file_type=ext.lstrip("."))


async def parse_upload(file: UploadFile, max_bytes: int) -> ParsedDocument:
    if not file.filename:
        raise HTTPException(status_code=400, detail="파일명이 필요합니다.")

    chunks: list[bytes] = []
    total_bytes = 0
    while True:
        chunk = await file.read(READ_CHUNK_BYTES)
        if not chunk:
            break
        total_bytes += len(chunk)
        if total_bytes > max_bytes:
            raise HTTPException(
                status_code=413,
                detail=f"파일 크기는 {max_bytes // (1024 * 1024)}MB 이하여야 합니다.",
            )
        chunks.append(chunk)
    data = b"".join(chunks)
    if len(data) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"파일 크기는 {max_bytes // (1024 * 1024)}MB 이하여야 합니다.",
        )
    if len(data) == 0:
        raise HTTPException(status_code=422, detail="빈 파일입니다.")

    try:
        return parse_bytes(file.filename, data, file.content_type)
    except HTTPException:
        raise
    except Exception:
        logger.exception("Document parsing failed for upload: filename=%s", file.filename)
        raise HTTPException(
            status_code=422,
            detail="문서 파싱에 실패했습니다. 파일 형식과 내용을 확인해 주세요.",
        ) from None
