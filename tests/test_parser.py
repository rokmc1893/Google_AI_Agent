"""
test_parser.py  ── 모듈 A: 비정형 계약서 파싱 모듈 TDD
═══════════════════════════════════════════════════════════════════════════════

[RED 단계] 구현 코드보다 먼저 작성된 테스트 코드입니다.

테스트 시나리오:
  ✅ 성공 케이스
    - PDF 파싱 후 계층형 딕셔너리(조/항/호) 반환 검증
    - Word(.docx) 파싱 후 동일한 계층형 구조 반환 검증
    - 다중 조항(조/항/호가 모두 존재) 파싱 정확도 검증
    - 조항 텍스트 무결성(원본 내용 보존) 검증

  ❌ 예외 케이스
    - 존재하지 않는 파일 경로 입력 시 FileNotFoundError
    - 지원하지 않는 파일 형식(.txt 등) 입력 시 UnsupportedFormatError
    - 빈 파일(0 bytes) 입력 시 EmptyDocumentError
    - 계약서 구조가 없는(조항 없는) 텍스트 파싱 시 빈 리스트 반환
    - 손상된 PDF 입력 시 CorruptedFileError

실행 방법:
  pytest tests/test_parser.py -v
  pytest tests/test_parser.py -v --cov=modules.parser
"""

import io
import json
import os
import textwrap
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

# 구현 모듈 import (Red 단계: 아직 미구현이므로 ImportError 예상)
from modules.parser import (
    ContractParser,
    CorruptedFileError,
    EmptyDocumentError,
    ParsedContract,
    UnsupportedFormatError,
    parse_contract_file,
)


# ─────────────────────────────────────────────────────────────────────────────
# 헬퍼: 테스트용 더미 파일 생성
# ─────────────────────────────────────────────────────────────────────────────
@pytest.fixture(scope="module")
def sample_docx_path(tmp_path_factory, sample_contract_text) -> Path:
    """샘플 계약서를 .docx 파일로 생성."""
    from docx import Document

    tmp = tmp_path_factory.mktemp("docs")
    path = tmp / "sample_contract.docx"
    doc = Document()
    for line in sample_contract_text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(path))
    return path


@pytest.fixture(scope="module")
def sample_pdf_path(tmp_path_factory, sample_contract_text) -> Path:
    """샘플 계약서를 .pdf 파일로 생성 (fpdf2 활용)."""
    pytest.importorskip("fpdf", reason="fpdf2 필요: pip install fpdf2")
    from fpdf import FPDF

    tmp = tmp_path_factory.mktemp("docs")
    path = tmp / "sample_contract.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    for line in sample_contract_text.splitlines():
        # ASCII 범위로 제한 (한글 폰트 없는 환경 대비 mock으로 대체 가능)
        safe_line = line.encode("ascii", errors="replace").decode("ascii")
        pdf.cell(0, 6, safe_line, new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(path))
    return path


@pytest.fixture()
def empty_pdf_path(tmp_path) -> Path:
    """0바이트 빈 파일."""
    path = tmp_path / "empty.pdf"
    path.write_bytes(b"")
    return path


@pytest.fixture()
def corrupted_pdf_path(tmp_path) -> Path:
    """손상된 PDF (헤더만 존재)."""
    path = tmp_path / "corrupted.pdf"
    path.write_bytes(b"%PDF-1.4 %%broken content%%")
    return path


@pytest.fixture()
def txt_file_path(tmp_path) -> Path:
    """지원하지 않는 형식(.txt)."""
    path = tmp_path / "contract.txt"
    path.write_text("제1조 테스트", encoding="utf-8")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# [GREEN] 성공 케이스 테스트
# ─────────────────────────────────────────────────────────────────────────────
class TestParserSuccessCases:
    """파싱 성공 시나리오 검증."""

    def test_parse_docx_returns_parsed_contract(
        self, sample_docx_path: Path
    ):
        """Word 파일 파싱 결과가 ParsedContract 타입을 반환해야 한다."""
        result = parse_contract_file(sample_docx_path)
        assert isinstance(result, ParsedContract)

    def test_parse_docx_has_articles(self, sample_docx_path: Path):
        """파싱 결과에 최소 1개 이상의 조(Article)가 포함되어야 한다."""
        result = parse_contract_file(sample_docx_path)
        assert len(result.articles) >= 1, "조항이 하나 이상 추출되어야 합니다."

    def test_parse_docx_article_structure(self, sample_docx_path: Path):
        """각 조(Article)는 title, number, paragraphs 키를 가져야 한다."""
        result = parse_contract_file(sample_docx_path)
        for article in result.articles:
            assert "number" in article, f"'number' 키 누락: {article}"
            assert "title" in article, f"'title' 키 누락: {article}"
            assert "paragraphs" in article, f"'paragraphs' 키 누락: {article}"

    def test_parse_docx_paragraph_structure(self, sample_docx_path: Path):
        """각 항(Paragraph)은 number, text, sub_items 키를 가져야 한다."""
        result = parse_contract_file(sample_docx_path)
        # 제2조 이상이 존재하고, 항이 있는 경우에만 검증
        articles_with_paragraphs = [
            a for a in result.articles if a.get("paragraphs")
        ]
        assert articles_with_paragraphs, "항이 있는 조가 없습니다."
        para = articles_with_paragraphs[0]["paragraphs"][0]
        assert "number" in para
        assert "text" in para
        assert "sub_items" in para

    def test_parse_docx_correct_article_count(self, sample_docx_path: Path):
        """샘플 계약서에서 정확히 4개의 조를 추출해야 한다."""
        result = parse_contract_file(sample_docx_path)
        assert len(result.articles) == 4, (
            f"예상 조항 수: 4, 실제: {len(result.articles)}"
        )

    def test_parse_docx_article_title_content(self, sample_docx_path: Path):
        """제1조의 제목은 '목적'을 포함해야 한다."""
        result = parse_contract_file(sample_docx_path)
        first_article = result.articles[0]
        assert "목적" in first_article["title"], (
            f"제1조 제목 오류: {first_article['title']}"
        )

    def test_parse_docx_sub_items_extracted(self, sample_docx_path: Path):
        """제2조 제1항 하위에 호(sub_item)가 추출되어야 한다."""
        result = parse_contract_file(sample_docx_path)
        article_2 = next(
            (a for a in result.articles if a["number"] == 2), None
        )
        assert article_2 is not None, "제2조를 찾지 못했습니다."
        para_1 = next(
            (p for p in article_2["paragraphs"] if p["number"] == 1), None
        )
        assert para_1 is not None, "제2조 제1항을 찾지 못했습니다."
        assert len(para_1["sub_items"]) >= 1, "호(sub_item)가 추출되지 않았습니다."

    def test_parse_returns_json_serializable(self, sample_docx_path: Path):
        """파싱 결과를 JSON으로 직렬화할 수 있어야 한다."""
        result = parse_contract_file(sample_docx_path)
        try:
            serialized = json.dumps(result.to_dict(), ensure_ascii=False)
            assert len(serialized) > 0
        except (TypeError, ValueError) as e:
            pytest.fail(f"JSON 직렬화 실패: {e}")

    def test_parse_preserves_text_integrity(
        self, sample_docx_path: Path, sample_contract_text: str
    ):
        """파싱된 텍스트를 이어붙이면 원본 핵심 키워드가 포함되어야 한다."""
        result = parse_contract_file(sample_docx_path)
        full_text = result.to_plain_text()
        # 핵심 법률 키워드 보존 검증
        for keyword in ["손해배상", "위약금", "비밀유지", "대금"]:
            assert keyword in full_text, f"키워드 '{keyword}'가 파싱 결과에 없습니다."

    def test_parse_pdf_with_mock(self, tmp_path: Path, sample_contract_text: str):
        """PDF 파싱을 pdfplumber mock으로 검증 (한글 폰트 의존성 제거)."""
        pdf_path = tmp_path / "mock_contract.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 mock")  # 더미 바이트

        mock_page = MagicMock()
        mock_page.extract_text.return_value = sample_contract_text

        with patch("pdfplumber.open") as mock_open:
            mock_open.return_value.__enter__.return_value.pages = [mock_page]
            result = parse_contract_file(pdf_path)

        assert isinstance(result, ParsedContract)
        assert len(result.articles) >= 1

    def test_parser_class_instantiation(self):
        """ContractParser 클래스를 인스턴스화할 수 있어야 한다."""
        parser = ContractParser()
        assert parser is not None

    def test_parser_parse_text_directly(self, sample_contract_text: str):
        """ContractParser.parse_text()는 텍스트를 직접 파싱할 수 있어야 한다."""
        parser = ContractParser()
        result = parser.parse_text(sample_contract_text)
        assert isinstance(result, ParsedContract)
        assert len(result.articles) == 4


# ─────────────────────────────────────────────────────────────────────────────
# [RED → GREEN] 예외 케이스 테스트
# ─────────────────────────────────────────────────────────────────────────────
class TestParserEdgeCases:
    """파싱 예외 및 엣지케이스 검증."""

    def test_file_not_found_raises_error(self, tmp_path: Path):
        """존재하지 않는 파일 경로 입력 시 FileNotFoundError를 발생시켜야 한다."""
        non_existent = tmp_path / "not_exist.pdf"
        with pytest.raises(FileNotFoundError):
            parse_contract_file(non_existent)

    def test_unsupported_format_raises_error(self, txt_file_path: Path):
        """지원하지 않는 파일 형식(.txt) 입력 시 UnsupportedFormatError를 발생시켜야 한다."""
        with pytest.raises(UnsupportedFormatError):
            parse_contract_file(txt_file_path)

    def test_empty_file_raises_error(self, empty_pdf_path: Path):
        """0바이트 빈 파일 입력 시 EmptyDocumentError를 발생시켜야 한다."""
        with pytest.raises(EmptyDocumentError):
            parse_contract_file(empty_pdf_path)

    def test_no_article_structure_returns_empty(self, tmp_path: Path):
        """조항 구조가 없는 텍스트 파싱 시 articles가 빈 리스트여야 한다."""
        docx_path = tmp_path / "no_structure.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("이것은 계약서 형식이 아닌 일반 텍스트입니다.")
        doc.save(str(docx_path))

        result = parse_contract_file(docx_path)
        assert result.articles == [], (
            f"구조 없는 문서의 articles는 빈 리스트여야 합니다: {result.articles}"
        )

    def test_corrupted_pdf_raises_error(self, corrupted_pdf_path: Path):
        """손상된 PDF 입력 시 CorruptedFileError를 발생시켜야 한다."""
        with pytest.raises(CorruptedFileError):
            parse_contract_file(corrupted_pdf_path)

    def test_parse_text_with_empty_string(self):
        """빈 문자열 입력 시 articles가 빈 리스트인 ParsedContract를 반환해야 한다."""
        parser = ContractParser()
        result = parser.parse_text("")
        assert isinstance(result, ParsedContract)
        assert result.articles == []

    def test_parse_text_with_whitespace_only(self):
        """공백만 있는 문자열 입력 시 articles가 빈 리스트여야 한다."""
        parser = ContractParser()
        result = parser.parse_text("   \n\t  ")
        assert result.articles == []

    def test_string_path_is_accepted(self, sample_docx_path: Path):
        """Path 객체 대신 문자열 경로를 입력해도 동작해야 한다."""
        result = parse_contract_file(str(sample_docx_path))
        assert isinstance(result, ParsedContract)

    def test_article_number_is_integer(self, sample_docx_path: Path):
        """조(Article)의 number 값은 정수(int)여야 한다."""
        result = parse_contract_file(sample_docx_path)
        for article in result.articles:
            assert isinstance(article["number"], int), (
                f"number가 int가 아닙니다: {type(article['number'])}"
            )

    def test_paragraphs_number_is_integer(self, sample_docx_path: Path):
        """항(Paragraph)의 number 값은 정수(int)여야 한다."""
        result = parse_contract_file(sample_docx_path)
        for article in result.articles:
            for para in article.get("paragraphs", []):
                assert isinstance(para["number"], int), (
                    f"paragraph number가 int가 아닙니다: {type(para['number'])}"
                )
