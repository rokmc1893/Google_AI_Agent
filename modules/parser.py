"""
parser.py  ── 모듈 A: 비정형 계약서 파싱 모듈
═══════════════════════════════════════════════════════════════════════════════

[Green 단계] 테스트를 통과하는 최소 구현 → [Refactor] 성능/가독성 개선

주요 기능:
  - PDF / Word(.docx) 파일을 입력받아 조/항/호 계층형 구조로 파싱
  - ParsedContract 데이터 클래스로 결과 반환
  - 비정상 파일(빈 파일, 손상 파일, 미지원 형식) 에러 처리

한국 법률 계약서 구조:
  제N조 (제목)          ← Article (조)
    제N항 본문 내용.     ← Paragraph (항)
      제N호 세부 사항.   ← SubItem (호)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Union


# ─────────────────────────────────────────────────────────────────────────────
# 커스텀 예외 클래스
# ─────────────────────────────────────────────────────────────────────────────
class UnsupportedFormatError(ValueError):
    """지원하지 않는 파일 형식 오류."""
    pass


class EmptyDocumentError(ValueError):
    """빈 문서 오류."""
    pass


class CorruptedFileError(IOError):
    """손상된 파일 오류."""
    pass


# ─────────────────────────────────────────────────────────────────────────────
# 지원 포맷 및 정규식 패턴
# ─────────────────────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

# 제1조 (목적) 또는 제1조(목적) 형식 모두 지원
_ARTICLE_PATTERN = re.compile(
    r"^제\s*(\d+)\s*조\s*[\(（](.+?)[\)）]",
    re.MULTILINE,
)
# 제1항 또는 ① 형식
_PARAGRAPH_PATTERN = re.compile(
    r"^[\s　]*제\s*(\d+)\s*항\s+",
    re.MULTILINE,
)
# 제1호
_SUBITEM_PATTERN = re.compile(
    r"^[\s　]*제\s*(\d+)\s*호\s+",
    re.MULTILINE,
)


# ─────────────────────────────────────────────────────────────────────────────
# 데이터 클래스
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class SubItem:
    """호(sub-item) 단위."""
    number: int
    text: str

    def to_dict(self) -> dict[str, Any]:
        return {"number": self.number, "text": self.text}


@dataclass
class Paragraph:
    """항(paragraph) 단위."""
    number: int
    text: str
    sub_items: list[SubItem] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "number": self.number,
            "text": self.text,
            "sub_items": [s.to_dict() for s in self.sub_items],
        }


@dataclass
class Article:
    """조(article) 단위."""
    number: int
    title: str
    paragraphs: list[Paragraph] = field(default_factory=list)
    raw_text: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "number": self.number,
            "title": self.title,
            "paragraphs": [p.to_dict() for p in self.paragraphs],
            "raw_text": self.raw_text,
        }


@dataclass
class ParsedContract:
    """파싱된 계약서 전체 결과."""
    source_path: str
    articles: list[dict[str, Any]]  # Article.to_dict() 형태의 딕셔너리 리스트
    raw_text: str = ""
    total_articles: int = 0

    def __post_init__(self):
        self.total_articles = len(self.articles)

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_path": self.source_path,
            "total_articles": self.total_articles,
            "articles": self.articles,
        }

    def to_plain_text(self) -> str:
        """파싱된 모든 텍스트를 이어붙인 평문 반환."""
        return self.raw_text


# ─────────────────────────────────────────────────────────────────────────────
# 핵심 파서 클래스 (Refactor: 책임 분리)
# ─────────────────────────────────────────────────────────────────────────────
class ContractParser:
    """
    비정형 계약서(PDF/Word)를 계층형 구조로 파싱하는 핵심 클래스.

    [Refactor] 각 메서드는 단일 책임 원칙(SRP)을 따르도록 분리되었습니다.
    """

    # ── 텍스트 파싱 ──────────────────────────────────────────────────────────
    def parse_text(self, text: str) -> ParsedContract:
        """
        평문 텍스트를 계층형 ParsedContract로 변환합니다.

        Args:
            text: 계약서 원문 텍스트

        Returns:
            ParsedContract: 파싱된 계약서 구조체
        """
        if not text or not text.strip():
            return ParsedContract(source_path="", articles=[], raw_text=text)

        articles = self._extract_articles(text)
        return ParsedContract(
            source_path="",
            articles=[a.to_dict() for a in articles],
            raw_text=text,
        )

    # ── 파일 파싱 ────────────────────────────────────────────────────────────
    def parse_file(self, file_path: Union[str, Path]) -> ParsedContract:
        """
        파일 경로를 받아 형식을 판별 후 파싱합니다.

        Args:
            file_path: PDF 또는 Word 파일 경로

        Returns:
            ParsedContract: 파싱된 계약서 구조체

        Raises:
            FileNotFoundError: 파일이 존재하지 않을 때
            UnsupportedFormatError: 지원하지 않는 형식일 때
            EmptyDocumentError: 파일이 비어 있을 때
            CorruptedFileError: 파일이 손상되었을 때
        """
        path = Path(file_path)

        # 파일 존재 여부
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {path}")

        # 지원 형식 검사
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(
                f"지원하지 않는 파일 형식: {path.suffix}. "
                f"지원 형식: {SUPPORTED_EXTENSIONS}"
            )

        # 빈 파일 검사
        if path.stat().st_size == 0:
            raise EmptyDocumentError(f"빈 파일입니다: {path}")

        # 형식별 텍스트 추출
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = self._extract_text_from_pdf(path)
        elif suffix in (".docx", ".doc"):
            text = self._extract_text_from_docx(path)
        else:
            raise UnsupportedFormatError(f"지원하지 않는 형식: {suffix}")

        articles = self._extract_articles(text)
        return ParsedContract(
            source_path=str(path),
            articles=[a.to_dict() for a in articles],
            raw_text=text,
        )

    # ── PDF 텍스트 추출 ──────────────────────────────────────────────────────
    def _extract_text_from_pdf(self, path: Path) -> str:
        """
        pdfplumber를 사용하여 PDF에서 텍스트를 추출합니다.

        [Refactor] 페이지별 스트리밍 방식으로 메모리 최적화
        """
        try:
            # pyrefly: ignore [missing-import]
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber를 설치하세요: pip install pdfplumber")

        try:
            pages_text: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                if not pdf.pages:
                    raise EmptyDocumentError(f"PDF에 페이지가 없습니다: {path}")
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)

            full_text = "\n".join(pages_text)
            if not full_text.strip():
                raise EmptyDocumentError(f"PDF에서 텍스트를 추출할 수 없습니다: {path}")
            return full_text

        except EmptyDocumentError:
            raise
        except Exception as e:
            # pdfplumber가 파싱 불가능한 파일에서 예외를 던지면 CorruptedFileError로 래핑
            if "invalid" in str(e).lower() or "corrupt" in str(e).lower():
                raise CorruptedFileError(f"손상된 PDF 파일입니다: {path}") from e
            # 기타 예외도 CorruptedFileError로 처리
            raise CorruptedFileError(f"PDF 파싱 중 오류 발생: {path} — {e}") from e

    # ── Word 텍스트 추출 ─────────────────────────────────────────────────────
    def _extract_text_from_docx(self, path: Path) -> str:
        """
        python-docx를 사용하여 Word 파일에서 텍스트를 추출합니다.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx를 설치하세요: pip install python-docx")

        try:
            doc = Document(str(path))
            lines = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(lines)
        except Exception as e:
            raise CorruptedFileError(
                f"Word 파일 파싱 중 오류 발생: {path} — {e}"
            ) from e

    # ── 계층 구조 추출 ───────────────────────────────────────────────────────
    def _extract_articles(self, text: str) -> list[Article]:
        """
        텍스트에서 조/항/호 계층 구조를 추출합니다.

        [Refactor] 조 단위로 텍스트를 분리 후 항, 호 순서로 파싱
        """
        articles: list[Article] = []

        # 조(Article) 경계로 텍스트 분리
        splits = _ARTICLE_PATTERN.split(text)
        # split() 결과: [이전 텍스트, 번호, 제목, 본문, 번호, 제목, 본문, ...]
        # 인덱스: 0=선행부, 1,2,3=첫번째 조, 4,5,6=두번째 조, ...
        if len(splits) < 4:
            # 조항 구조 없음
            return []

        idx = 1  # 첫 번째 조부터
        while idx + 2 <= len(splits):
            try:
                art_num = int(splits[idx])
                art_title = splits[idx + 1].strip()
                art_body = splits[idx + 2] if idx + 2 < len(splits) else ""
            except (ValueError, IndexError):
                idx += 3
                continue

            paragraphs = self._extract_paragraphs(art_body)
            article = Article(
                number=art_num,
                title=art_title,
                paragraphs=paragraphs,
                raw_text=art_body.strip(),
            )
            articles.append(article)
            idx += 3

        return articles

    def _extract_paragraphs(self, text: str) -> list[Paragraph]:
        """항(Paragraph) 추출."""
        paragraphs: list[Paragraph] = []
        splits = _PARAGRAPH_PATTERN.split(text)
        # split() 결과: [이전, 번호, 본문, 번호, 본문, ...]
        if len(splits) < 3:
            return []

        idx = 1
        while idx + 1 <= len(splits) - 1:
            try:
                para_num = int(splits[idx])
                para_body = splits[idx + 1]
            except (ValueError, IndexError):
                idx += 2
                continue

            # 다음 항 또는 조 시작 전까지의 텍스트만 추출
            sub_items = self._extract_sub_items(para_body)
            paragraph = Paragraph(
                number=para_num,
                text=para_body.strip().splitlines()[0] if para_body.strip() else "",
                sub_items=sub_items,
            )
            paragraphs.append(paragraph)
            idx += 2

        return paragraphs

    def _extract_sub_items(self, text: str) -> list[SubItem]:
        """호(SubItem) 추출."""
        sub_items: list[SubItem] = []
        splits = _SUBITEM_PATTERN.split(text)
        if len(splits) < 3:
            return []

        idx = 1
        while idx + 1 <= len(splits) - 1:
            try:
                item_num = int(splits[idx])
                item_text = splits[idx + 1].strip().splitlines()[0]
            except (ValueError, IndexError):
                idx += 2
                continue
            sub_items.append(SubItem(number=item_num, text=item_text))
            idx += 2

        return sub_items


# ─────────────────────────────────────────────────────────────────────────────
# 편의 함수 (공개 API)
# ─────────────────────────────────────────────────────────────────────────────
def parse_contract_file(file_path: Union[str, Path]) -> ParsedContract:
    """
    계약서 파일을 파싱하는 최상위 편의 함수.

    Args:
        file_path: PDF 또는 Word 파일 경로 (str 또는 Path)

    Returns:
        ParsedContract: 파싱된 계약서 계층 구조

    Examples:
        >>> result = parse_contract_file("contract.docx")
        >>> result.articles[0]["title"]
        '목적'
    """
    parser = ContractParser()
    return parser.parse_file(file_path)
